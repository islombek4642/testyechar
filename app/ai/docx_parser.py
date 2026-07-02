"""
DOCXQuizParser — Classic-format DOCX reader for the AI Resolver.
Extracts ?/+/= paragraphs and inline images per question.
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from typing import List

from docx import Document
from docx.oxml.ns import qn

from app.models.question import RawOption, RawQuestion
from app.utils.logger import get_logger

log = get_logger(__name__)

_DRAWING = qn("w:drawing")
_PICT = qn("w:pict")
_BLIP = qn("a:blip")
_R_EMBED = qn("r:embed")


@dataclass
class AIRawQuestion:
    """RawQuestion with an optional attached image (bytes)."""

    question: RawQuestion
    image: bytes | None = None


class DOCXQuizParser:
    """
    Parses a DOCX file in Classic quiz format (?/+/=) into AIRawQuestion objects.
    Inline images that appear between a question paragraph and its options are
    attached to the preceding AIRawQuestion.

    Note: only ``doc.paragraphs`` is traversed — table cells are not visited.
    Table-layout DOCX files are out of scope for this parser; use the
    pipeline's TableParser for those.
    """

    def parse(self, content: bytes) -> List[AIRawQuestion]:
        try:
            doc = Document(BytesIO(content))
        except Exception as exc:
            log.error("DOCX faylni ochishda xatolik: %s", exc)
            return []

        results: List[AIRawQuestion] = []
        current_rq: RawQuestion | None = None
        current_img: bytes | None = None
        line_number = 0

        for para in doc.paragraphs:
            line_number += 1
            text = para.text.strip()

            # Detect image paragraph (has drawing/pict element, no text marker)
            has_drawing = para._element.find(".//" + _DRAWING) is not None
            has_pict = para._element.find(".//" + _PICT) is not None
            is_image_para = (has_drawing or has_pict) and not text

            if is_image_para:
                if current_rq is not None and current_img is None:
                    current_img = self._extract_image(para, doc)
                continue

            if not text:
                continue

            if text.startswith("!"):
                continue

            marker = text[0]
            body = text[1:].strip()

            if marker == "?":
                self._flush(results, current_rq, current_img)
                current_rq = RawQuestion(question_text=body, line_start=line_number)
                current_img = None

            elif marker == "+":
                if current_rq is None:
                    current_rq = RawQuestion(question_text="", line_start=line_number)
                current_rq.options.append(RawOption(text=body, is_correct=True))

            elif marker == "=":
                if current_rq is None:
                    current_rq = RawQuestion(question_text="", line_start=line_number)
                current_rq.options.append(RawOption(text=body, is_correct=False))

            else:
                if current_rq is not None:
                    if current_rq.options:
                        last = current_rq.options[-1]
                        last.text = (last.text + " " + text).strip()
                    else:
                        current_rq.question_text = (
                            current_rq.question_text + " " + text
                        ).strip()

        self._flush(results, current_rq, current_img)
        return results

    @staticmethod
    def _flush(
        results: List[AIRawQuestion],
        rq: RawQuestion | None,
        img: bytes | None,
    ) -> None:
        if rq is not None and (rq.question_text.strip() or rq.options):
            results.append(AIRawQuestion(question=rq, image=img))

    @staticmethod
    def _extract_image(para, doc) -> bytes | None:
        """Extract the first inline image from a paragraph via its blip relationship."""
        try:
            for blip in para._element.iter(_BLIP):
                r_id = blip.get(_R_EMBED)
                if r_id:
                    part = doc.part.related_parts.get(r_id)
                    if part is not None:
                        return part.blob
        except Exception as exc:
            log.warning("Rasmni olishda xatolik: %s", exc)
        return None
