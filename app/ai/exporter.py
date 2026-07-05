"""
AI Result Exporter.

Exports merged resolved questions to:
- Compact JSON  ({"q":..,"o":[..],"c":N,"cf":X})
- Full JSON     (includes source, tier metadata)
- DOCX          (classic ? / + / = format)
- TXT           (classic ? / + / = format)
"""

from __future__ import annotations

import json
from pathlib import Path
from typing import List

from app.ai.merger import ResolvedQuestion
from app.ai.statistics import AIStatistics
from app.utils.logger import get_logger

log = get_logger(__name__)


class AIExporter:
    """Exports AI-resolved questions to multiple formats."""

    # ── JSON ──────────────────────────────────────────────────────────────

    def to_compact_json_string(self, questions: List[ResolvedQuestion]) -> str:
        """Compact JSON: q, o, c, cf only."""
        data = [
            {"q": q.q, "o": q.o, "c": q.c, "cf": q.cf}
            for q in questions
        ]
        return json.dumps(data, ensure_ascii=False)

    def to_full_json_string(self, questions: List[ResolvedQuestion]) -> str:
        """Full JSON: includes source, tier, cf for every question."""
        data = [q.model_dump() for q in questions]
        return json.dumps(data, ensure_ascii=False, indent=2)

    def to_compact_json_file(self, questions: List[ResolvedQuestion], path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(self.to_compact_json_string(questions), encoding="utf-8")
        log.info(f"Compact JSON eksport qilindi: {path}")
        return path

    def to_full_json_file(self, questions: List[ResolvedQuestion], path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(self.to_full_json_string(questions), encoding="utf-8")
        log.info(f"Full JSON eksport qilindi: {path}")
        return path

    # ── TXT (Classic ? / + / =) ───────────────────────────────────────────

    def to_txt_string(self, questions: List[ResolvedQuestion]) -> str:
        """Classic TXT format: ? / + / = with blank lines between questions."""
        lines: List[str] = []
        for q in questions:
            lines.append(f"? {q.q}")
            for idx, opt in enumerate(q.o):
                marker = "+" if idx == q.c else "="
                lines.append(f"{marker} {opt}")
            lines.append("")  # blank separator
        return "\n".join(lines)

    def to_txt_file(self, questions: List[ResolvedQuestion], path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(self.to_txt_string(questions), encoding="utf-8")
        log.info(f"TXT eksport qilindi: {path}")
        return path

    # ── DOCX ──────────────────────────────────────────────────────────────

    def to_docx_bytes(self, questions: List[ResolvedQuestion]) -> bytes:
        """
        Generate a DOCX file in-memory and return raw bytes.
        Uses python-docx with clean monospace-style formatting.
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError("python-docx kutubxonasi o'rnatilmagan. 'pip install python-docx' buyrug'ini ishga tushiring.")

        import io

        doc = Document()

        # Document title style
        title = doc.add_heading("Quiz — AI Resolved", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # spacer

        for i, q in enumerate(questions, start=1):
            # Question paragraph — Classic ? marker, no number prefix
            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"? {q.q}")
            q_run.bold = True
            q_run.font.size = Pt(11)

            # Options
            for idx, opt in enumerate(q.o):
                marker = "+" if idx == q.c else "="
                opt_para = doc.add_paragraph()
                opt_run = opt_para.add_run(f"{marker} {opt}")
                opt_run.font.size = Pt(10.5)
                if marker == "+":
                    opt_run.font.color.rgb = RGBColor(0x16, 0xa3, 0x4a)  # green
                else:
                    opt_run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)  # slate

            doc.add_paragraph()  # blank line between questions

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()

    def to_docx_file(self, questions: List[ResolvedQuestion], path: Path) -> Path:
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_bytes(self.to_docx_bytes(questions))
        log.info(f"DOCX eksport qilindi: {path}")
        return path

    # ── Statistics Report (TXT) ────────────────────────────────────────────

    def to_stats_report_string(self, stats: AIStatistics) -> str:
        """Human-readable statistics report in Uzbek."""
        lines = [
            "=" * 55,
            "   AI JAVOB TOPUVCHI — STATISTIKA HISOBOTI",
            "=" * 55,
            f"  Model               : {stats.model}",
            f"  Rejim               : {'Batch API' if stats.used_batch_api else 'Standart API'}",
            f"  Jami savollar       : {stats.total_questions}",
            f"  Avvaldan yechilgan  : {stats.already_resolved}",
            f"  AI ga yuborilgan    : {stats.sent_to_ai}",
            f"  AI yechgan          : {stats.ai_resolved}",
            f"  Muvaffaqiyatsiz     : {stats.failed}",
            "",
            "  — Ishonch darajasi —",
            f"  Ishonchli (≥0.90)   : {stats.trusted_count}",
            f"  Ogohlantirish       : {stats.warning_count}",
            f"  Qo'lda tekshirish   : {stats.manual_review_count}",
            f"  O'rtacha ishonch    : {stats.average_confidence:.3f}",
            "",
            "  — Tokenlar —",
            f"  Kiritish tokenlar   : {stats.input_tokens:,}",
            f"  Chiqish tokenlar    : {stats.output_tokens:,}",
            f"  Jami tokenlar       : {stats.total_tokens:,}",
            "",
            "  — Narx —",
            f"  Standart API narxi  : ${stats.cost_standard_usd:.4f}",
            f"  Batch API narxi     : ${stats.cost_batch_usd:.4f}",
            "",
            f"  Chunk soni          : {stats.chunk_count}",
            f"  Ishlash vaqti       : {stats.duration_seconds}s",
            "=" * 55,
        ]
        if stats.warnings:
            lines.append("")
            lines.append("  OGOHLANTIRISHLAR:")
            for w in stats.warnings:
                lines.append(f"  ⚠  {w}")
            lines.append("=" * 55)
        return "\n".join(lines)
