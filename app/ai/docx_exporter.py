"""
AIDocxExporter — reconstruct a Classic-format DOCX from AI resolver results.
Images are preserved; correct answers are marked with '+'.
"""
from __future__ import annotations

from io import BytesIO
from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Inches

from app.utils.logger import get_logger

if TYPE_CHECKING:
    from app.ai.merger import MergeResult

log = get_logger(__name__)


class AIDocxExporter:
    """Exports a MergeResult to a Classic-format DOCX file (bytes)."""

    def export(self, merge_result: "MergeResult", filename: str = "result.docx") -> bytes:
        """
        Build and return a DOCX where each question is written in Classic format:
          ? question text
          [image — if merge_result.images[i] exists]
          + correct option   (omitted if c == -1)
          = wrong options
        """
        doc = Document()

        for i, q in enumerate(merge_result.questions):
            doc.add_paragraph(f"? {q.q}")

            # Insert image if available for this question
            img_bytes = merge_result.images.get(i)
            if img_bytes:
                try:
                    doc.add_picture(BytesIO(img_bytes), width=Inches(3))
                except Exception as exc:
                    log.warning("Rasmni DOCX ga yozishda xatolik (savol %d): %s", i, exc)

            # Write options
            for opt_idx, opt_text in enumerate(q.o):
                if opt_idx == q.c:
                    doc.add_paragraph(f"+ {opt_text}")
                else:
                    doc.add_paragraph(f"= {opt_text}")

            # Blank separator
            doc.add_paragraph("")

        buf = BytesIO()
        doc.save(buf)
        log.info("DOCX eksport tayyor: %d ta savol, fayl=%s", len(merge_result.questions), filename)
        return buf.getvalue()
