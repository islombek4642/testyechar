"""
JSON exporter.

Serialises validated questions to compact JSON or DOCX, ready for download.
"""

from __future__ import annotations

import io
import json
import re
import unicodedata
from pathlib import Path
from typing import List

from app.config import settings
from app.models.question import ParsedQuestion
from app.utils.logger import get_logger

log = get_logger(__name__)

# ── Math-italic → ASCII normalisation map ──────────────────────────────
# Characters in U+1D400–U+1D7FF ("Mathematical Alphanumeric Symbols") are
# not covered by common fonts (Calibri, Arial, Times New Roman) so Word
# renders them as "?".  We map them back to plain ASCII before writing DOCX.
# The mapping is built once at import time using unicodedata.

# Math italic Greek letter names → standard Unicode Greek chars.
# Covers both small (α β γ …) and capital (Α Β Γ …) variants.
# Unicode uses "LAMDA" (not LAMBDA) — included under both spellings.
_GREEK_NAME_TO_CHAR: dict[str, tuple[str, str]] = {
    "ALPHA":   ("α", "Α"), "BETA":    ("β", "Β"), "GAMMA":   ("γ", "Γ"),
    "DELTA":   ("δ", "Δ"), "EPSILON": ("ε", "Ε"), "ZETA":    ("ζ", "Ζ"),
    "ETA":     ("η", "Η"), "THETA":   ("θ", "Θ"), "IOTA":    ("ι", "Ι"),
    "KAPPA":   ("κ", "Κ"), "LAMDA":   ("λ", "Λ"), "LAMBDA":  ("λ", "Λ"),
    "MU":      ("μ", "Μ"), "NU":      ("ν", "Ν"), "XI":      ("ξ", "Ξ"),
    "OMICRON": ("ο", "Ο"), "PI":      ("π", "Π"), "RHO":     ("ρ", "Ρ"),
    "SIGMA":   ("σ", "Σ"), "TAU":     ("τ", "Τ"), "UPSILON": ("υ", "Υ"),
    "PHI":     ("φ", "Φ"), "CHI":     ("χ", "Χ"), "PSI":     ("ψ", "Ψ"),
    "OMEGA":   ("ω", "Ω"),
    # variant symbols
    "EPSILON SYMBOL": ("ϵ", ""),  "THETA SYMBOL": ("ϑ", ""),
    "KAPPA SYMBOL":   ("ϰ", ""),  "PHI SYMBOL":   ("ϕ", ""),
    "RHO SYMBOL":     ("ϱ", ""),  "PI SYMBOL":    ("ϖ", ""),
}
# Pre-compile word-boundary patterns once
_GREEK_PATTERNS = [
    (re.compile(r"\b" + name + r"\b"), small, cap)
    for name, (small, cap) in _GREEK_NAME_TO_CHAR.items()
]


def _build_math_latin_map() -> dict[str, str]:
    result: dict[str, str] = {}
    for cp in range(0x1D400, 0x1D800):
        try:
            ch = chr(cp)
            name = unicodedata.name(ch, "")
            if not name:
                continue

            # Latin letters: "MATHEMATICAL ITALIC SMALL X" → "x"
            hit = re.search(r"\b(SMALL|CAPITAL)\s+([A-Z])\b", name)
            if hit:
                letter = hit.group(2)
                result[ch] = letter.lower() if hit.group(1) == "SMALL" else letter
                continue

            # Greek letters: "MATHEMATICAL ITALIC SMALL ALPHA" → "α"
            for pat, small, cap in _GREEK_PATTERNS:
                if pat.search(name):
                    if "SMALL" in name and small:
                        result[ch] = small
                    elif "CAPITAL" in name and cap:
                        result[ch] = cap
                    break

        except Exception:
            pass

    # Extra letterlike symbols used in math typography
    result.update({
        "ℎ": "h",  # PLANCK CONSTANT (italic h)
        "ℯ": "e",  # SCRIPT SMALL E
        "ℐ": "I",  # SCRIPT CAPITAL I
        "ℓ": "l",  # SCRIPT SMALL L
        "ℕ": "N",  # DOUBLE-STRUCK CAPITAL N
        "ℙ": "P",  # DOUBLE-STRUCK CAPITAL P
        "ℚ": "Q",  # DOUBLE-STRUCK CAPITAL Q
        "ℝ": "R",  # DOUBLE-STRUCK CAPITAL R
        "ℤ": "Z",  # DOUBLE-STRUCK CAPITAL Z
        "⁡": "",  # FUNCTION APPLICATION (invisible separator)
    })
    return result


_MATH_LATIN_MAP: dict[str, str] = _build_math_latin_map()
_MATH_LATIN_TABLE = str.maketrans(_MATH_LATIN_MAP)


def _docx_safe(text: str) -> str:
    """
    Convert Mathematical Alphanumeric Symbol letters to plain ASCII so they
    render correctly in Word without requiring specialised math fonts.

    Greek letters (α β π …) and math operators (√ ∞ ∑ …) are kept as-is
    because they are covered by standard fonts (Cambria, Arial Unicode MS).
    """
    return text.translate(_MATH_LATIN_TABLE)


class JSONExporter:
    """
    Export validated questions to JSON or DOCX format.

    Supports:
    - Compact (single-line) output for production use
    - Pretty-printed output for debugging
    - Writing to file or returning as string
    - DOCX with Unicode-safe rendering
    """

    def to_json_string(
        self,
        questions: List[ParsedQuestion],
        indent: int | None = None,
    ) -> str:
        data = [q.model_dump() for q in questions]
        return json.dumps(data, ensure_ascii=False, indent=indent)

    def to_file(
        self,
        questions: List[ParsedQuestion],
        output_path: Path | None = None,
        indent: int | None = None,
    ) -> Path:
        if output_path is None:
            settings.ensure_dirs()
            output_path = settings.output_dir / settings.output_filename

        if indent is None:
            indent = settings.json_indent

        json_str = self.to_json_string(questions, indent=indent)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        output_path.write_text(json_str, encoding="utf-8")

        log.info(
            f"Exported [bold green]{len(questions)}[/] questions to "
            f"[cyan]{output_path}[/]"
        )
        return output_path

    def to_docx_bytes(self, questions: List[ParsedQuestion]) -> bytes:
        """
        Generate a DOCX file in-memory and return raw bytes, in the same
        plain Classic-format style as AIDocxExporter (app/ai/docx_exporter.py)
        so the Parser's export and the AI-resolved export look identical:
          ? question text
          [image — if any]
          + correct option
          = wrong options

        Mathematical italic Unicode letters (𝑥, 𝑙𝑜𝑔, 𝑐𝑡𝑔…) are still
        converted to plain ASCII so they display correctly in Word without
        requiring specialised math fonts; this is a correctness fix, not
        styling, so it stays regardless of the plain layout.
        """
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError:
            raise ImportError(
                "python-docx kutubxonasi o'rnatilmagan. "
                "'pip install python-docx' buyrug'ini ishga tushiring."
            )

        doc = Document()

        for q in questions:
            doc.add_paragraph(f"? {_docx_safe(q.q)}")

            for img_rel_path in q.img:
                img_path = settings.output_dir / img_rel_path
                if img_path.exists():
                    try:
                        doc.add_picture(str(img_path), width=Inches(3))
                    except Exception as exc:
                        log.warning(f"Rasmni DOCX ga qo'shishda xatolik: {exc}")

            for opt_idx, opt in enumerate(q.o):
                marker = "+" if opt_idx == q.c else "="
                doc.add_paragraph(f"{marker} {_docx_safe(opt)}")

            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf.read()
