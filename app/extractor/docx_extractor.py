"""
python-docx based DOCX text extractor with highlight and color detection.

Extracts plain text from DOCX files paragraph-by-paragraph, maintaining unicode.
Automatically detects option highlighting (yellow or colored background) and 
font colors (green text) in Word to convert options to correct answers (+ / #).
"""

from __future__ import annotations

import io
import re
from pathlib import Path
from typing import Any, Generator

from app.config import settings
from app.parser.image_tokens import build_image_token
from app.utils.logger import get_logger
from .base import BaseExtractor, ExtractionResult

log = get_logger(__name__)


class DOCXExtractor(BaseExtractor):
    """
    Concrete DOCX extractor using python-docx.

    Features:
    - Structurally ordered paragraph extraction (handles both body paragraphs and tables)
    - Highlight detection (yellow or any background fill)
    - Font color detection (specifically green text indicating correct answers)
    - Never crashes — all errors are caught and logged
    """

    @property
    def name(self) -> str:
        return "python-docx"

    def supports(self, path: Path) -> bool:
        """Return True if this extractor can handle the given file extension."""
        return path.suffix.lower() == ".docx"

    def extract(self, path: Path) -> ExtractionResult:
        """
        Open *path* and extract text structurally, detecting highlighted/green options.

        Returns:
            ExtractionResult with a single page representing the full document content.
        """
        log.info(f"[bold cyan]Opening DOCX[/]: {path.name}")

        try:
            from docx import Document
            from docx.text.paragraph import Paragraph
        except ImportError:
            log.error("python-docx is not installed.")
            return ExtractionResult(
                pages=[],
                page_count=0,
                source_path=path,
                extractor_name=self.name,
            )

        try:
            doc = Document(str(path))
        except Exception as exc:
            log.error(f"Error opening DOCX: {exc}")
            return ExtractionResult(
                pages=[],
                page_count=0,
                source_path=path,
                extractor_name=self.name,
            )

        extracted_lines: list[str] = []

        # --- Inline-paragraph format detection (e.g. "Question\na. opt\nb. opt" in one paragraph) ---
        # Structure: each paragraph contains one complete question with options embedded via
        # soft line-breaks (Shift+Enter). Options use lowercase "a. b. c. d." notation.
        # No correct answer is marked — all options exported as "= " (unknown).
        try:
            para_texts = [
                self._extract_paragraph_text(p)
                for p in self._iter_all_paragraphs(doc)
            ]
            if self._is_inline_paragraph_format(para_texts):
                log.info("DOCX: inline-paragraph formati aniqlandi (a. b. variant) — klassik formatga o'girilmoqda.")
                extracted_lines = self._extract_inline_paragraphs(para_texts)
                full_text = "\n".join(extracted_lines)
                return ExtractionResult(
                    pages=[full_text],
                    page_count=1,
                    source_path=path,
                    extractor_name=self.name,
                )
        except Exception as exc:
            log.warning(f"Inline-paragraph detection failed, fallback to standard: {exc}")

        # --- Jadval-Blok format detection (tables with ==== / # / ++++ separators) ---
        # Structure: many single-column tables (one per question).
        # Row 1: question text. Even rows: ==== separators. Odd rows: options (# = correct).
        # Body paragraphs between tables contain ++++ question separators.
        try:
            if self._is_docx_table_blocks_format(doc):
                log.info("DOCX: jadval-blok formati aniqlandi — blok formatga o'girilmoqda.")
                extracted_lines = self._extract_docx_table_blocks(doc)
                full_text = "\n".join(extracted_lines)
                return ExtractionResult(
                    pages=[full_text],
                    page_count=1,
                    source_path=path,
                    extractor_name=self.name,
                )
        except Exception as exc:
            log.warning(f"Jadval-blok detection failed, fallback to standard: {exc}")

        # --- Table-Bold format detection (e.g. anatomy test banks) ---
        # Structure: single-column table, bold rows = questions, normal rows = options.
        # First option after each question is treated as the correct answer.
        try:
            if self._is_table_bold_format(doc):
                log.info("DOCX: jadval+bold formati aniqlandi — klassik formatga o'girilmoqda.")
                extracted_lines = self._extract_table_bold(doc)
                full_text = "\n".join(extracted_lines)
                return ExtractionResult(
                    pages=[full_text],
                    page_count=1,
                    source_path=path,
                    extractor_name=self.name,
                )
        except Exception as exc:
            log.warning(f"Table-bold detection failed, fallback to standard: {exc}")

        # --- Multi-column table Q&A format ---
        # Structure: each table row = one question.
        # col[0]=question, '*'-prefixed col=correct, rest=wrong answers.
        # Layout B variant: col[0]=sequential number, col[1]=question, col[2+]=answers.
        # NOTE: must run before numbered_bold_option_format, which can false-positive
        # on answer cells like "3 ta" (number+Uzbek) or "S. surname" (letter+period).
        try:
            if self._is_table_columns_qa_format(doc):
                log.info("DOCX: ko'p ustunli jadval Q&A formati aniqlandi.")
                extracted_lines = self._extract_table_columns_qa(doc, path)
                full_text = "\n".join(extracted_lines)
                return ExtractionResult(
                    pages=[full_text],
                    page_count=1,
                    source_path=path,
                    extractor_name=self.name,
                )
        except Exception as exc:
            log.warning(f"Table-columns Q&A detection failed, fallback: {exc}")

        # --- Numbered + Letter + Bold Correct Answer format ---
        # Structure: body paragraphs. Questions: "1. text", "2. text".
        # Options: single Cyrillic or Latin letter + "." or ")" + text.
        # Correct answer = option whose content (beyond letter prefix) is bold.
        try:
            _para_objs_nb = list(self._iter_all_paragraphs(doc))
            _para_texts_nb = [self._extract_paragraph_text(p) for p in _para_objs_nb]
            if self._is_numbered_bold_option_format(_para_texts_nb, _para_objs_nb):
                log.info("DOCX: raqamli+harfli+bold formati aniqlandi — klassik formatga o'girilmoqda.")
                extracted_lines = self._extract_numbered_bold_options(_para_texts_nb, _para_objs_nb)
                full_text = "\n".join(extracted_lines)
                return ExtractionResult(
                    pages=[full_text],
                    page_count=1,
                    source_path=path,
                    extractor_name=self.name,
                )
        except Exception as exc:
            log.warning(f"Numbered-bold option detection failed, fallback: {exc}")

        # Iterate structurally through paragraphs in both body and tables
        try:
            # --- Pass 1: extract raw text + remember paragraph objects ---
            paragraphs: list[tuple[str, Any]] = []
            for p in self._iter_all_paragraphs(doc):
                paragraphs.append((self._extract_paragraph_text(p), p))

            p_count = len(paragraphs)

            # --- Priority check: does the document already use explicit markers? ---
            # If so, the author marked correct answers via text (+ / #a)) — trust
            # that and ignore highlighting entirely (prevents decorative colour
            # from being misread as a second "correct" indicator, which would
            # make the validator see "2 correct answers" for one question).
            # Paragraphs can bundle a question and all its options into one
            # multi-line blob (soft line-breaks), so the marker may sit mid-blob —
            # split on "\n" first so each option lands on its own line for the check.
            para_lines: list[str] = []
            for text, _ in paragraphs:
                para_lines.extend(text.split("\n"))
            has_explicit_correct = self.has_explicit_correct(para_lines)

            if has_explicit_correct:
                log.info(
                    f"DOCX tahlil qilindi: {p_count} ta paragraf. "
                    "Hujjatda aniq belgi topildi — belgilar ishonchli manba, "
                    "bo'yash tekshiruvi o'tkazib yuborildi."
                )
                extracted_lines = [text for text, _ in paragraphs]

            else:
                # --- Pass 2: no explicit markers → use highlighting ---
                highlight_count = 0
                numbering_count = 0
                bold_fallback_count = 0

                # Group into blocks: each starts at a paragraph that looks
                # like a new question (Word list-numbered, or already
                # marker-prefixed) and includes every paragraph after it up
                # to the next such start as its options -- so correctness
                # can be judged across a whole question's options at once
                # instead of one paragraph in isolation. A block starts at
                # a paragraph carrying Word list numbering (numPr) or an
                # already-present "?" marker -- NOT at "-"/"–"/"="/"+"/"#",
                # which mark OPTIONS and must stay inside the current
                # block. Some documents give their options no marker of
                # any kind at all, relying purely on color/bold to single
                # out the answer; keying block starts on any of those
                # option markers would split each such option into its
                # own bogus single-paragraph "question".
                blocks: list[list[tuple[str, Any]]] = []
                current_block: list[tuple[str, Any]] = []
                for p_text, p in paragraphs:
                    stripped = p_text.strip()
                    is_new_block_start = bool(stripped) and (
                        stripped.startswith("?") or self._has_list_numbering(p)
                    )
                    if is_new_block_start:
                        if current_block:
                            blocks.append(current_block)
                        current_block = [(p_text, p)]
                    elif stripped:
                        current_block.append((p_text, p))
                    # blank paragraphs are dropped, matching prior behavior
                    # (they never carried a "-"/"–" prefix either)
                if current_block:
                    blocks.append(current_block)

                for block in blocks:
                    q_text, q_p = block[0]
                    option_entries = block[1:]

                    stripped_q = q_text.strip()
                    is_numbered_q = (
                        stripped_q
                        and stripped_q[0] not in "?=+#-–"
                        and self._has_list_numbering(q_p)
                    )
                    if is_numbered_q:
                        # Word auto-numbered ("1.", "2." ...) question --
                        # the number is list metadata, not run text, so
                        # it's otherwise indistinguishable from ordinary
                        # prose.
                        extracted_lines.append("? " + stripped_q)
                        numbering_count += 1
                    else:
                        extracted_lines.append(q_text)

                    # A STRONG signal (highlight, dominant red/green) on
                    # any option settles it outright -- first one wins,
                    # matching prior behavior.
                    strong_idx = None
                    for i, (_, o_p) in enumerate(option_entries):
                        if self._check_is_correct_option(o_p):
                            strong_idx = i
                            break

                    if strong_idx is not None:
                        for i, (o_text, _) in enumerate(option_entries):
                            if i == strong_idx:
                                processed = self._process_highlighted_paragraph_text(o_text)
                                if processed != o_text:
                                    highlight_count += 1
                                extracted_lines.append(processed)
                            else:
                                extracted_lines.append(o_text)
                    else:
                        # No highlight/red/green anywhere in this block --
                        # fall back to relative bold: exactly one option
                        # whose every run is bold, regardless of shared
                        # color, distinguishes it from its non-bold
                        # siblings. Covers both a plain-black-bold scheme
                        # AND a document that colors an entire question
                        # block uniformly (e.g. all dark blue) and relies
                        # on bold alone within that block for the answer.
                        bold_flags = [self._is_fully_bold(o_p) for _, o_p in option_entries]
                        if bold_flags.count(True) == 1:
                            bold_idx = bold_flags.index(True)
                            for i, (o_text, _) in enumerate(option_entries):
                                if i == bold_idx:
                                    processed = self._process_highlighted_paragraph_text(o_text)
                                    if processed != o_text:
                                        bold_fallback_count += 1
                                    extracted_lines.append(processed)
                                else:
                                    extracted_lines.append(o_text)
                        else:
                            for o_text, _ in option_entries:
                                extracted_lines.append(o_text)

                log.info(
                    f"DOCX tahlil qilindi: {p_count} ta paragraf topildi, "
                    f"{highlight_count} ta variantda rang/bo'yash orqali, "
                    f"{bold_fallback_count} ta variantda nisbiy qalinlik orqali to'g'rilik belgisi aniqlandi, "
                    f"{numbering_count} ta savolga ro'yxat raqamlashi asosida '?' qo'yildi."
                )
                # Convert "1. Savol" → "? Savol" only when no explicit markers exist,
                # because ABCParser needs numbered questions intact for format detection.
                extracted_lines = self._mark_numbered_questions(extracted_lines)
                # Some documents number neither the question nor mark it any
                # other way -- the question paragraph is just plain text
                # (often itself ending in "?", but with no marker at the
                # START of the line) immediately followed by its "-"
                # (dash-prefixed) options. Without a "?" prefix, ClassicParser
                # can't tell the question apart from ordinary unrecognized
                # text and silently drops it, leaving its first option to
                # open a bogus question with empty text.
                extracted_lines = self._mark_unmarked_dash_questions(extracted_lines)

        except Exception as exc:
            log.error(f"Error during DOCX extraction loop: {exc}", exc_info=True)

        extracted_lines = self.fix_all_correct(extracted_lines)
        full_text = "\n".join(extracted_lines)
        full_text = self._split_inline_markers(full_text)
        full_text = "\n".join(
            self._number_unnumbered_abc_questions(full_text.split("\n"))
        )

        result = ExtractionResult(
            pages=[full_text],
            page_count=1,
            source_path=path,
            extractor_name=self.name,
        )

        log.info(
            f"Extraction complete — "
            f"{result.character_count:,} characters extracted from DOCX"
        )

        return result

    # ── OMML (Word equation) extraction ────────────────────────────────────

    _W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    _A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"
    _R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

    @staticmethod
    def _run_text(run_elem) -> list[str]:
        """Collect w:t and w:br (soft return) children of a single w:r run."""
        parts: list[str] = []
        for elem in run_elem:
            elem_local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if elem_local == "t":
                parts.append(elem.text or "")
            elif elem_local == "br":
                parts.append("\n")
        return parts

    def _extract_paragraph_text(self, paragraph) -> str:
        """
        Extract text from a paragraph including OMML math equations and
        hyperlink-wrapped runs.

        python-docx's paragraph.text skips <m:oMath> elements entirely.
        This method walks the paragraph XML and converts math elements to
        readable plain text (e.g. m:f fraction → "π/6"). It also descends
        into <w:hyperlink> wrapper elements -- Word auto-links things like
        bare email addresses and URLs, nesting their run(s) one level
        deeper than ordinary text; skipping that wrapper (as a plain
        direct-children walk would) silently drops the linked text
        entirely, which is disastrous when it's the correct answer itself.
        """
        parts: list[str] = []
        for child in paragraph._element:
            local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if local == "r":
                parts.extend(self._run_text(child))
            elif local == "hyperlink":
                for grandchild in child:
                    gc_local = (
                        grandchild.tag.split("}")[-1]
                        if "}" in grandchild.tag
                        else grandchild.tag
                    )
                    if gc_local == "r":
                        parts.extend(self._run_text(grandchild))
            elif local in ("oMath", "oMathPara"):
                parts.append(self._omml_to_text(child))
        return "".join(parts)

    def _omml_to_text(self, elem) -> str:
        """Recursively convert an OMML element tree to plain text."""
        local = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        if local == "f":
            # Fraction: <m:num> / <m:den>
            num = elem.find(f"{{{self._M_NS}}}num")
            den = elem.find(f"{{{self._M_NS}}}den")
            num_text = self._omml_to_text(num) if num is not None else ""
            den_text = self._omml_to_text(den) if den is not None else ""
            return f"{num_text}/{den_text}"

        if local == "r":
            # Math run — collect m:t text
            t = elem.find(f"{{{self._M_NS}}}t")
            return (t.text or "") if t is not None else ""

        if local == "sSup":
            # Superscript: base^exp
            base = elem.find(f"{{{self._M_NS}}}e")
            sup = elem.find(f"{{{self._M_NS}}}sup")
            base_text = self._omml_to_text(base) if base is not None else ""
            sup_text = self._omml_to_text(sup) if sup is not None else ""
            return f"{base_text}^{sup_text}"

        if local == "sSub":
            # Subscript: base_sub
            base = elem.find(f"{{{self._M_NS}}}e")
            sub = elem.find(f"{{{self._M_NS}}}sub")
            base_text = self._omml_to_text(base) if base is not None else ""
            sub_text = self._omml_to_text(sub) if sub is not None else ""
            return f"{base_text}_{sub_text}"

        if local == "rad":
            # Radical (square root or nth root)
            deg = elem.find(f"{{{self._M_NS}}}deg")
            base = elem.find(f"{{{self._M_NS}}}e")
            deg_text = self._omml_to_text(deg) if deg is not None else ""
            base_text = self._omml_to_text(base) if base is not None else ""
            if deg_text.strip():
                return f"{deg_text}√{base_text}"
            return f"√{base_text}"

        if local == "d":
            # Delimiter (parentheses/brackets)
            parts: list[str] = []
            for e in elem.findall(f"{{{self._M_NS}}}e"):
                parts.append(self._omml_to_text(e))
            return "(" + ", ".join(parts) + ")"

        # Skip property nodes (fPr, rPr, etc.) and recurse into content nodes
        parts = []
        for child in elem:
            child_local = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if child_local.endswith("Pr"):
                continue
            parts.append(self._omml_to_text(child))
        return "".join(parts)

    def _iter_all_paragraphs(self, doc) -> Generator[Any, None, None]:
        """
        Iterate structurally through paragraphs in document body, including tables.
        Preserves original document layout order.
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        for element in doc.element.body:
            tag = element.tag
            if tag.endswith('p'):
                yield Paragraph(element, doc)
            elif tag.endswith('tbl'):
                table = Table(element, doc)
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph

    def _check_is_correct_option(self, paragraph) -> bool:
        """
        Scan all runs in the paragraph for a STRONG correct-answer signal:
        highlighting, or dominant green/red text color. Deliberately does
        NOT look at bold here -- bold alone is ambiguous per-paragraph (see
        _is_fully_bold and the block-level fallback in the Pass 2 loop,
        which only trusts bold when nothing in the whole question carries
        one of these strong signals).
        """
        for run in paragraph.runs:
            # 1. Highlight detection
            try:
                if run.font.highlight_color is not None:
                    return True
            except (ValueError, AttributeError):
                # python-docx throws ValueError for unknown highlight values (e.g. "none")
                pass

            # 2. Dominant green or red text detection -- some test banks mark
            # the correct option in green, others in red (often paired with
            # bold), against a neutral/blue color for the wrong options.
            if run.font.color and run.font.color.rgb:
                try:
                    r, g, b = run.font.color.rgb
                    is_dominant_green = g > r + 40 and g > b + 40
                    is_dominant_red = r > g + 40 and r > b + 40
                    if is_dominant_green or is_dominant_red:
                        return True
                except Exception:
                    pass

        return False

    @staticmethod
    def _is_fully_bold(paragraph) -> bool:
        """True if the paragraph has text and every non-empty run is bold,
        regardless of color."""
        has_any_run_text = False
        for run in paragraph.runs:
            if run.text.strip():
                has_any_run_text = True
                if not run.bold:
                    return False
        return has_any_run_text

    @staticmethod
    def _has_list_numbering(paragraph) -> bool:
        """
        True if the paragraph is part of a Word auto-numbered list (numPr).

        Word renders the "1.", "2." etc. visually from list-numbering
        metadata, not as literal characters in the run text -- python-docx
        extracts such a question as plain text with no leading digit at
        all. Checking numPr directly is the reliable structural signal for
        "this paragraph is a question" in documents laid out this way,
        rather than inferring it from what text happens to follow.
        """
        from docx.oxml.ns import qn

        pPr = paragraph._p.find(qn("w:pPr"))
        return pPr is not None and pPr.find(qn("w:numPr")) is not None

    _NUMBERED_Q = re.compile(
        r"^\d+[\.\)]\s+(.+)$"              # group 1: "1. Q"  "1) Q"
        r"|^\d{2,3}\s+[\.\)]+\s*(.+)$"    # group 2: "172 .Q"  "17 .Q"
        r"|^\d{2,3}\s+\d+[\.\)]\s*(.+)$"  # group 3: "291 3.Q"
        r"|^\d{2,3}([А-ЯЁA-Z].{4,})$"    # group 4: "119Kompetensiya"
        r"|^\d{2,3}\s+([А-ЯЁA-Z].{4,})$", # group 5: "185 Ta‟lim"
        re.UNICODE
    )

    def _mark_numbered_questions(self, lines: list[str]) -> list[str]:
        """
        Convert numbered question lines to Classic ? format.

        "1. Savol matni"    →  "? Savol matni"
        "2) Savol matni"    →  "? Savol matni"
        "172 .Jamiyat..."   →  "? Jamiyat..."
        "119Kompetensiya"   →  "? Kompetensiya"
        "185 Ta‟limning..." →  "? Ta‟limning..."

        Lines that already start with a Classic/ABC marker (?, =, +, #) are
        left untouched, as are blank lines and short lines without content.
        """
        result = []
        for line in lines:
            stripped = line.strip()
            if stripped and stripped[0] not in "?=+#":
                m = self._NUMBERED_Q.match(stripped)
                if m:
                    q_text = next((g for g in m.groups() if g is not None), "")
                    line = "? " + q_text
            result.append(line)
        return result

    def _mark_unmarked_dash_questions(self, lines: list[str]) -> list[str]:
        """
        Give a "?" marker to a question line that has neither a number nor
        any marker of its own -- it's just plain text (often itself ending
        in "?", but that's punctuation, not a leading marker) -- when the
        next non-blank line is a "-"/"–" (hyphen or en-dash) option. That's
        the one structural signal that survives extraction for this format:
        a question is always immediately followed by its first option.
        Covers documents whose questions have no Word list numbering
        either (see _has_list_numbering, which handles the numbered case).
        """
        result = list(lines)
        n = len(result)
        for i, line in enumerate(result):
            stripped = line.strip()
            if not stripped or stripped[0] in "?=+#-–":
                continue
            j = i + 1
            while j < n and not result[j].strip():
                j += 1
            # The next option may already have been converted to "+text"
            # (correct-marker, dash stripped) by the highlighting pass
            # above, rather than a bare "-"/"–", if it's the very first
            # (and correct) option.
            if j < n and result[j].strip().startswith(("-", "–", "+")):
                result[i] = "? " + stripped
        return result

    # An ABC option's first item: "a)" / "#a)" / "A." — used to locate where a
    # question block starts when the question line itself carries no number.
    _OPTION_A_START = re.compile(r"^#?\s*[aA][\.\)]\s*\S")
    _ANY_OPTION_START = re.compile(r"^#?\s*[a-fA-F][\.\)]")
    _NUMBERED_LINE = re.compile(r"^\d+[\.\)]\s")

    def _number_unnumbered_abc_questions(self, lines: list[str]) -> list[str]:
        """
        Give a sequential number to ABC question lines that have none of
        their own, so ABCParser can tell where each question begins.

        Word's auto-numbered lists (numPr) render numbers visually but don't
        store them in the run text — python-docx extracts such questions as
        plain text with no leading digit. Without a number (or a "?" marker),
        ABCParser can't distinguish "new question" from "continuation of the
        previous option", and the whole file collapses into one giant question.

        The one reliable signal that survives extraction is structural: a
        question is always immediately followed by its "a)" option. Any
        marker-less, number-less line directly preceding an "a)" option gets
        "N. " prepended (sequential, independent of any numbers already
        present elsewhere — ABCParser only needs *a* leading number).

        Only activates when "a)" option starts clearly outnumber already
        numbered questions, i.e. the document mixes numbered/un-numbered
        questions or has none numbered — leaving fully-numbered ABC files
        (where counts roughly match) untouched.
        """
        a_starts = sum(1 for ln in lines if self._OPTION_A_START.match(ln.strip()))
        numbered = sum(1 for ln in lines if self._NUMBERED_LINE.match(ln.strip()))
        if a_starts < 5 or a_starts <= numbered:
            return lines

        result = list(lines)
        counter = 0
        for i, line in enumerate(result):
            stripped = line.strip()
            if (
                not stripped
                or stripped[0] in "?+=#"
                or self._ANY_OPTION_START.match(stripped)
                or self._NUMBERED_LINE.match(stripped)
            ):
                continue

            j = i + 1
            while j < len(result) and not result[j].strip():
                j += 1
            if j < len(result) and self._OPTION_A_START.match(result[j].strip()):
                lead = len(line) - len(line.lstrip())
                counter += 1
                result[i] = line[:lead] + f"{counter}. " + stripped

        if counter:
            log.info(
                f"DOCX: {counter} ta raqamsiz savolga ketma-ket raqam berildi "
                "(ABC format chegaralarini aniqlash uchun)."
            )
        return result

    # Two-pass inline option splitting:
    # Pass 1 — uppercase A-D: split after lowercase letter or punctuation.
    #   Excluding uppercase A-Z and digits from the lookbehind prevents splitting
    #   inside abbreviations like "(LMA)" where M precedes A).
    # Pass 2 — lowercase a-d: split ONLY after punctuation (not after any letter/digit)
    #   so that "chiqishga)" or "(1938-yilda)" are NOT incorrectly split.
    # Both also exclude a preceding "(" — "(a)"/"(b)" is a fully-bracketed
    # parenthetical label (e.g. "Ikkinchi(a) va uchinchi(b) ...", referring to
    # sub-parts named a/b), not the start of a new option; a real option
    # marker is never itself wrapped in its own opening paren.
    # Also exclude a preceding "-"/"–" (hyphen/en-dash): that's the Classic
    # wrong-answer marker itself, not the end of a preceding option's text --
    # "- A) a-1, b-2" is ONE option whose own content happens to read like a
    # matching-pairs answer, not "-" (an empty option) glued to a fresh "A)"
    # option that needs splitting off.
    _INLINE_OPT_UPPER = re.compile(
        r'(?<=[^\s#*+=\-–\nA-Z\d(])'  # preceded by lowercase letter or punctuation (not uppercase/digit/open-paren/dash)
        r'(\s*)'
        r'([#*+=]?[A-D]\))',      # uppercase A-D option marker
        re.UNICODE,
    )
    _INLINE_OPT_LOWER = re.compile(
        r'(?<=[^\s#*+=\-–\na-zA-Z\d(])'  # preceded by punctuation only (not open-paren/dash)
        r'(\s*)'
        r'([#*+=]?[a-d]\))',          # lowercase a-d option marker
        re.UNICODE,
    )
    # Matches a 2-3 digit question number mid-line (e.g. "...text 99. Next...")
    _INLINE_QNUM = re.compile(
        r'(?<=[^\n\d])'        # preceded by non-digit, non-newline
        r'(\s+)'               # at least one space before number
        r'(\d{2,3}[\.\)](?=\s))',  # 2-3 digit number + . or )
        re.UNICODE,
    )

    # Classic markers (= / +) that appear after a semicolon mid-line.
    # e.g. "a) wrong;= also wrong;*d) correct" → each gets its own line.
    _INLINE_CLASSIC = re.compile(r';\s*([=+])\s', re.UNICODE)

    # Normalize "b )" / "c )" (letter + space + bracket) → "b)" / "c)"
    # Some DOCX files have a stray space between the option letter and bracket,
    # and may also have leading whitespace from Word's visual indentation.
    _OPT_SPACE_BRACKET = re.compile(r'^\s*([a-fA-F]) \)', re.MULTILINE | re.UNICODE)

    def _split_inline_markers(self, text: str) -> str:
        """
        Ensure every question number and option letter starts on its own line.
        Handles cases where Shift+Enter was not used and multiple items are on
        one line (e.g. "A) wrong;= wrong;*d) correct" or "A) wrong B) +C) correct").
        """
        # Normalize "b )" → "b)" (stray space between letter and bracket)
        text = self._OPT_SPACE_BRACKET.sub(r'\1)', text)
        # Split Classic markers that follow a semicolon (;= text or ;+ text)
        text = self._INLINE_CLASSIC.sub(lambda m: "\n" + m.group(1) + " ", text)
        # Split uppercase option markers (A-D) — safe after any non-ws char
        text = self._INLINE_OPT_UPPER.sub(lambda m: "\n" + m.group(2), text)
        # Split lowercase option markers (a-d) — only after punctuation
        text = self._INLINE_OPT_LOWER.sub(lambda m: "\n" + m.group(2), text)
        # Split at inline question numbers
        text = self._INLINE_QNUM.sub(lambda m: "\n" + m.group(2), text)
        return text

    def _process_highlighted_paragraph_text(self, p_text: str) -> str:
        """
        Format a highlighted option paragraph into the target correct format.
        """
        stripped = p_text.strip()
        if not stripped:
            return p_text

        # 1. Classic format option starting with '=', '-' or '–' (en-dash)
        if stripped[0] in "=-–":
            leading_spaces = len(p_text) - len(p_text.lstrip())
            content = p_text[leading_spaces + 1:]
            return p_text[:leading_spaces] + "+" + content

        # 2. ABC format option starting with 'A)' or 'B.'
        abc_match = re.match(r"^([a-fA-F][\.\)])\s*", stripped)
        if abc_match:
            if not stripped.startswith("#"):
                leading_spaces = len(p_text) - len(p_text.lstrip())
                return p_text[:leading_spaces] + "#" + p_text[leading_spaces:]

        # 3. Blocks format plain option
        is_question = (
            stripped.startswith("?") or 
            re.match(r"^\d+[\.\)]", stripped)
        )
        if not is_question and not stripped.startswith("#") and not stripped.startswith("+"):
            leading_spaces = len(p_text) - len(p_text.lstrip())
            return p_text[:leading_spaces] + "#" + p_text[leading_spaces:]

        return p_text

    # ── Inline-paragraph format detection & extraction ──────────────────
    # Format: each paragraph = one question with options embedded via soft newlines.
    # Options use "a. text", "b. text" notation (lowercase letter + period).

    _PARA_INLINE_OPT = re.compile(r'\n\s*[a-d]\.\s', re.IGNORECASE)
    # Matches options on the same line separated by 3+ spaces: "was   c. Would"
    _INLINE_OPT_DOT_SEP = re.compile(r'\s{3,}([a-d])\.\s*', re.IGNORECASE)
    _OPT_DOT_LINE = re.compile(r'^([a-d])\.\s*(.*)', re.IGNORECASE | re.DOTALL)
    _Q_NUMBER_PREFIX = re.compile(r'^\d+[\.\)]\s*(.*)', re.DOTALL)

    def _is_inline_paragraph_format(self, paragraph_texts: list[str]) -> bool:
        """
        Return True if paragraphs use the inline question+option format:
          - At least 5 paragraphs contain soft-newline embedded "a. b. c. d." options
          - No existing Classic markers (?, +, =) present
        """
        has_classic = any(
            t.strip() and t.strip()[0] in "?+="
            for t in paragraph_texts
        )
        if has_classic:
            return False
        inline_count = sum(
            1 for t in paragraph_texts
            if self._PARA_INLINE_OPT.search(t)
        )
        return inline_count >= 5

    def _extract_inline_paragraphs(self, paragraph_texts: list[str]) -> list[str]:
        """
        Convert inline-paragraph format to Classic marker lines.

        Paragraph with embedded options → "? question", "= opt_a", "= opt_b", ...
        No correct answer is marked (correct_index = -1 in downstream validator).
        """
        result: list[str] = []
        for text in paragraph_texts:
            if not self._PARA_INLINE_OPT.search(text):
                stripped = text.strip()
                if stripped:
                    result.append(stripped)
                continue

            # Normalize options sitting on the same line (3+ spaces between them)
            text = self._INLINE_OPT_DOT_SEP.sub(r'\n\1. ', text)

            parts = [p.strip() for p in text.split('\n') if p.strip()]
            if not parts:
                continue

            # First segment = question text (strip optional leading number "1. ")
            q_text = parts[0]
            q_num = self._Q_NUMBER_PREFIX.match(q_text)
            if q_num:
                q_text = q_num.group(1).strip()
            if q_text:
                result.append('? ' + q_text)

            # Remaining segments = options
            for part in parts[1:]:
                opt_m = self._OPT_DOT_LINE.match(part)
                if opt_m:
                    opt_text = opt_m.group(2).strip()
                    if opt_text:
                        result.append('= ' + opt_text)
                elif result:
                    # Continuation of the previous line
                    result[-1] += ' ' + part

        q_count = sum(1 for l in result if l.startswith('? '))
        log.info(
            f"Inline-paragraph extraction: {q_count} ta savol, "
            f"to'g'ri javob belgilanmagan (barcha variantlar '=' sifatida eksport qilinadi)."
        )
        return result

    # ── Jadval-Blok format detection & extraction ───────────────────────
    # Format: many single-column tables, one per question.
    # Row 1 = question text, alternating ==== separators and option rows.
    # Correct option rows start with '#'. Body paragraphs carry ++++ separators.

    def _is_docx_table_blocks_format(self, doc) -> bool:
        """
        Return True if the document uses the jadval-blok (multi-table blocks) format:
        - 5+ tables, each with a single column
        - Table rows contain ==== option separators and # correct markers
        """
        if len(doc.tables) < 5:
            return False

        sep_count = hash_count = 0
        sample = min(10, len(doc.tables))

        for tbl in doc.tables[:sample]:
            if len(tbl.columns) > 2:
                return False
            for row in tbl.rows:
                text = row.cells[0].text.strip() if row.cells else ""
                if len(text) >= 3 and all(c == "=" for c in text):
                    sep_count += 1
                elif text.startswith("#"):
                    hash_count += 1

        # Expect ≥3 separators per table and at least one # marker per 2 tables
        return sep_count >= sample * 3 and hash_count >= max(1, sample // 2)

    def _extract_docx_table_blocks(self, doc) -> list[str]:
        """
        Extract jadval-blok format into flat block-format lines for BlocksParser.

        Body <w:p> paragraphs → emitted as-is (++++ separators and header text).
        Body <w:tbl> tables   → each row's cell text (question / ==== / # answer / options).
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        lines: list[str] = []

        for element in doc.element.body:
            tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

            if tag == "p":
                para = Paragraph(element, doc)
                text = self._extract_paragraph_text(para).strip()
                if text:
                    lines.append(text)

            elif tag == "tbl":
                table = Table(element, doc)
                for row in table.rows:
                    if not row.cells:
                        continue
                    # Join all paragraphs inside the cell
                    cell_parts = [
                        self._extract_paragraph_text(p).strip()
                        for p in row.cells[0].paragraphs
                    ]
                    cell_text = " ".join(p for p in cell_parts if p)
                    if cell_text:
                        lines.append(cell_text)

        q_seps = sum(
            1 for l in lines
            if len(l.strip()) >= 3 and all(c == "+" for c in l.strip())
        )
        log.info(
            f"DOCX jadval-blok: {len(doc.tables)} jadval, "
            f"~{q_seps} ta savol ajratuvchisi topildi."
        )
        return lines

    # ── Table-Bold format detection & extraction ────────────────────────

    def _is_table_bold_format(self, doc) -> bool:
        """
        Return True if the document is a "table-bold" medical test bank:
        - At least one table with 1-2 columns
        - Bold paragraphs are questions, non-bold are options
        - Bold ratio 10–55% (for 4-option: 1 bold + 4 normal → 20%)
        - No existing Classic/ABC markers already in the text
        """
        if not doc.tables:
            return False
        tbl = doc.tables[0]
        if len(tbl.columns) > 2:
            return False

        bold_rows = 0
        total_rows = 0
        has_marker = False

        for row in tbl.rows[:60]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    total_rows += 1
                    if any(r.bold and r.text.strip() for r in para.runs):
                        bold_rows += 1
                    if text and text[0] in "?+=":
                        has_marker = True

        if total_rows < 10 or has_marker:
            return False
        ratio = bold_rows / total_rows
        return 0.10 <= ratio <= 0.55

    def _extract_table_bold(self, doc) -> list[str]:
        """
        Extract table-bold format into Classic marker lines.

        Bold paragraph  → "? question text"
        1st non-bold after question → "+ correct answer"
        Subsequent non-bold         → "= wrong answer"
        """
        lines: list[str] = []
        tbl = doc.tables[0]
        option_index = 0

        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = self._extract_paragraph_text(para).strip()
                    if not text:
                        continue

                    is_bold = any(r.bold and r.text.strip() for r in para.runs)

                    if is_bold:
                        lines.append("? " + text)
                        option_index = 0
                    else:
                        if option_index == 0:
                            lines.append("+ " + text)
                        else:
                            lines.append("= " + text)
                        option_index += 1

        q_count = sum(1 for l in lines if l.startswith("? "))
        log.info(
            f"Table-bold extraction: {q_count} ta savol, "
            f"birinchi variant to'g'ri javob sifatida belgilandi."
        )
        return lines

    # ── Numbered + Letter + Bold Correct Answer format ──────────────────
    # Format: body paragraphs with numbered questions and lettered options.
    # Any single letter (Cyrillic А-Я or Latin A-Z) + "." or ")" starts an option.
    # The option whose TEXT CONTENT (beyond the letter prefix) is bold = correct answer.
    # Supports compound paragraphs where multiple items share one paragraph via soft newlines.

    # Single letter (Cyrillic uppercase/lowercase + Ё/ё + Latin) + period/paren + non-space
    _LETTER_OPTION_RE = re.compile(
        r"^\s*[А-ЯЁа-яёA-Za-z][.)]\s*\S",
        re.UNICODE,
    )
    # Letter + period/paren + optional spaces — used to strip the prefix
    _OPTION_PREFIX_RE = re.compile(
        r"^[А-ЯЁа-яёA-Za-z][.)]\s*",
        re.UNICODE,
    )
    # A bare letter-marker prefix and nothing else — punctuation optional
    # ("А. ", "b)", "Д " — the last is the "loose" Cyrillic-only marker with
    # no period). Used to confirm a bold chunk's leading text is JUST the
    # option's letter marker (not part of the answer's actual content).
    _BOLD_PREFIX_RE = re.compile(
        r"^[А-ЯЁа-яёA-Za-z][.)]?\s*$",
        re.UNICODE,
    )
    # Flexible question: allows no space after period ("12.Text" or "12. Text")
    # Also handles "128 В каком..." (number + space + Cyrillic/Latin letter content)
    _NUMBERED_Q_FLEX = re.compile(
        r"^\s*\d+[.)]\s*(.*)"                    # 1. text  or  1) text
        r"|^\s*\d{1,3}\s+([А-ЯЁа-яёA-Za-z].*)",  # 128 В... (no punctuation after number)
        re.UNICODE,
    )

    def _is_numbered_bold_option_format(
        self, para_texts: list[str], para_objs: list
    ) -> bool:
        """
        Return True when document uses numbered questions + lettered options
        where correct answers are indicated by bold text content.

        Detection criteria:
        - No existing Classic markers (?, +, =)
        - ≥5 question blocks (numbered-text OR compound Q+options paragraphs —
          Word's auto-numbered lists don't store the visible "1." in para.text,
          so a compound paragraph whose first segment isn't itself an option,
          followed by ≥2 lettered-option segments, also counts as a question)
        - ≥5 letter-option occurrences (standalone paragraphs OR segments
          inside compound soft-newline-joined paragraphs)
        - ≥3 options whose content text is bold
        """
        has_classic = any(
            t.strip() and t.strip()[0] in "?+="
            for t in para_texts
        )
        if has_classic:
            return False

        q_count = 0
        opt_count = 0
        for t in para_texts:
            s = t.strip()
            if not s:
                continue
            if "\n" in s:
                segments = [seg.strip() for seg in s.split("\n") if seg.strip()]
                opt_segs = sum(1 for seg in segments if self._LETTER_OPTION_RE.match(seg))
                if opt_segs >= 2 and not self._LETTER_OPTION_RE.match(segments[0]):
                    q_count += 1
                opt_count += opt_segs
            elif self._LETTER_OPTION_RE.match(s):
                opt_count += 1
            elif self._NUMBERED_Q_FLEX.match(s):
                q_count += 1

        if q_count < 5 or opt_count < 5:
            return False

        bold_count = 0
        for text, para in zip(para_texts, para_objs):
            if self._paragraph_has_bold_option(text, para):
                bold_count += 1
                if bold_count >= 3:
                    return True
        return False

    def _get_bold_chunks(self, para) -> list[str]:
        """
        Merge consecutive bold runs into chunks and return their stripped texts.

        Word sometimes splits a single bold phrase across several runs (e.g.
        spell-check or formatting boundaries) — merging adjacent bold runs
        before comparing avoids missing matches due to that fragmentation.
        """
        chunks: list[str] = []
        current: list[str] = []
        for run in para.runs:
            if run.bold:
                current.append(run.text)
            else:
                if current:
                    chunks.append("".join(current))
                    current = []
        if current:
            chunks.append("".join(current))
        return [c.strip() for c in chunks if c.strip()]

    @staticmethod
    def _content_is_bold(opt_text: str, bold_chunks: list[str]) -> bool:
        """
        True when an option's CONTENT (letter prefix already stripped) is
        covered by one of the paragraph's bold chunks — either because only
        the content is bold ("b) " plain + "reward" bold → chunk == opt_norm),
        or because the whole segment including its letter prefix is bold
        ("А. текст" fully bold → chunk == "<marker><opt_norm>").

        Deliberately NOT a substring check: grammar-style options often
        share text ("pass" / "would pass" / "passed"), and a bold answer
        like "would have caught" would otherwise also "match" the shorter
        sibling options "caught", "would", "would have" — marking several
        options correct for one question. Only an exact match, or a match
        with a genuine letter-marker prefix ("А. "/"b) "), counts.
        """
        opt_norm = opt_text.strip()
        if not opt_norm:
            return False
        for chunk in bold_chunks:
            if opt_norm == chunk:
                return True
            if chunk.endswith(opt_norm):
                prefix = chunk[: -len(opt_norm)]
                if DOCXExtractor._BOLD_PREFIX_RE.fullmatch(prefix):
                    return True
        return False

    def _paragraph_has_bold_option(self, text: str, para) -> bool:
        """True when at least one lettered-option segment's content is bold."""
        bold_chunks = self._get_bold_chunks(para)
        if not bold_chunks:
            return False

        stripped = text.strip()
        segments = stripped.split("\n") if "\n" in stripped else [stripped]
        for seg in segments:
            seg_s = seg.strip()
            if self._LETTER_OPTION_RE.match(seg_s):
                opt_text = self._OPTION_PREFIX_RE.sub("", seg_s, count=1).strip()
                if self._content_is_bold(opt_text, bold_chunks):
                    return True
        return False

    def _extract_numbered_bold_options(
        self, para_texts: list[str], para_objs: list
    ) -> list[str]:
        """
        Convert numbered+letter+bold format to Classic marker lines.

        "1. Question" → "? Question"
        Bold-content option → "+ option text"
        Plain option → "= option text"

        Handles two paragraph structures:
        - Simple: one item per paragraph (standard case)
        - Compound: question + options or multiple options in one paragraph
          joined by soft newlines (\\n) — bold letters extracted from runs
        """
        result: list[str] = []
        in_sequence = False

        for text, para in zip(para_texts, para_objs):
            stripped = text.strip()
            if not stripped:
                continue

            # ── Compound paragraph: contains soft-newline separated segments ──
            if "\n" in stripped:
                segments = [s.strip() for s in stripped.split("\n") if s.strip()]
                has_opts = any(self._LETTER_OPTION_RE.match(s) for s in segments)

                if has_opts or any(self._NUMBERED_Q_FLEX.match(s) for s in segments):
                    bold_chunks = self._get_bold_chunks(para)
                    # Word's auto-numbered lists don't store the visible "1."
                    # in para.text — the question segment then carries no
                    # numeric prefix at all. The structural signal that still
                    # survives is paragraph position: segment 0 of a compound
                    # Q+options block is the question whenever the block has
                    # ≥2 lettered options, prefix or not.
                    opt_seg_count = sum(1 for s in segments if self._LETTER_OPTION_RE.match(s))
                    for idx, seg in enumerate(segments):
                        seg_s = seg.strip()
                        if not seg_s:
                            continue
                        q_m = self._NUMBERED_Q_FLEX.match(seg_s)
                        # Standard option: "А. text"  or  "А) text"
                        is_std_opt = bool(self._LETTER_OPTION_RE.match(seg))
                        # Loose option: "Д text" (single Cyrillic letter + space,
                        # no period). Latin letters are deliberately excluded —
                        # single-letter English words ("I decided...", "A man
                        # is...") would otherwise be misread as option markers.
                        loose_m = re.match(
                            r"^([А-ЯЁа-яё])\s+(\S.*)", seg_s, re.UNICODE
                        ) if not is_std_opt else None
                        is_unnumbered_q_start = (
                            idx == 0 and not is_std_opt and not loose_m and opt_seg_count >= 2
                        )

                        if (q_m and not is_std_opt and not loose_m) or is_unnumbered_q_start:
                            in_sequence = True
                            q_text = (q_m.group(1) or q_m.group(2)).strip() if q_m else seg_s
                            if q_text:
                                result.append("? " + q_text)
                        elif (is_std_opt or loose_m) and in_sequence:
                            if is_std_opt:
                                opt_text = self._OPTION_PREFIX_RE.sub("", seg_s, count=1).strip()
                            else:
                                opt_text = loose_m.group(2).strip()
                            if opt_text:
                                marker = "+" if self._content_is_bold(opt_text, bold_chunks) else "="
                                result.append(marker + " " + opt_text)
                        elif in_sequence and result and seg_s:
                            result[-1] += " " + seg_s
                    continue

            # ── Simple paragraph ──
            q_m = self._NUMBERED_Q_FLEX.match(stripped)
            is_opt = bool(self._LETTER_OPTION_RE.match(text))

            if q_m and not is_opt:
                in_sequence = True
                q_text = (q_m.group(1) or q_m.group(2) or "").strip()
                if q_text:
                    result.append("? " + q_text)
            elif is_opt and in_sequence:
                opt_text = self._OPTION_PREFIX_RE.sub("", stripped, count=1).strip()
                if opt_text:
                    bold_chunks = self._get_bold_chunks(para)
                    marker = "+" if self._content_is_bold(opt_text, bold_chunks) else "="
                    result.append(marker + " " + opt_text)
            elif in_sequence and result:
                result[-1] += " " + stripped

        q_count = sum(1 for l in result if l.startswith("? "))
        correct_count = sum(1 for l in result if l.startswith("+ "))
        log.info(
            f"Numbered-bold extraction: {q_count} ta savol, "
            f"{correct_count} ta to'g'ri javob bold orqali aniqlandi."
        )
        return result

    # ── Multi-column table Q&A format detection & extraction ────────────
    # Format: each table row = one Q&A set.
    # Layout A: col[0]=question, col with '*' prefix=correct, others=wrong.
    # Layout B: col[0]=sequential number, col[1]=question, col[2+]=answers.

    def _get_cell_text(self, cell) -> str:
        """
        Return a table cell's text, including OMML math equations.

        ``cell.text`` silently drops ``<m:oMath>`` content, so a cell whose
        only content is a Word equation (e.g. "∪", "A∩B") returns "" —
        making the cell look empty even though it holds an answer option.
        This reuses :meth:`_extract_paragraph_text` (already used for body
        paragraphs) for each paragraph in the cell.
        """
        return " ".join(
            self._extract_paragraph_text(p).replace('\n', ' ')
            for p in cell.paragraphs
        ).strip()

    def _get_unique_cell_texts(self, row) -> list[str]:
        """Return cell texts for a row, skipping horizontally-merged duplicate cells."""
        seen: set[int] = set()
        texts: list[str] = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id not in seen:
                seen.add(tc_id)
                texts.append(self._get_cell_text(cell))
        return texts

    _SANITIZE_STEM_RE = re.compile(r"[^A-Za-z0-9_-]+")

    def _sanitize_doc_stem(self, stem: str) -> str:
        """Sanitize a filename stem for use as an images subdirectory name."""
        cleaned = self._SANITIZE_STEM_RE.sub("_", stem).strip("_")
        return cleaned or "doc"

    def _extract_cell_images(
        self, doc, cell, doc_stem: str, table_idx: int, row_idx: int, counter: list[int]
    ) -> list[str]:
        """
        Extract embedded images from *cell*, save them under
        ``settings.images_dir / doc_stem``, and return their paths relative
        to ``settings.output_dir`` (e.g. ``images/<doc_stem>/t0_row5_img1.png``).

        EMF/WMF images are converted to PNG via Pillow; if conversion fails
        the original bytes are saved with their native extension instead.
        Never raises — extraction errors are logged and skipped.
        """
        paths: list[str] = []
        blip_tag = f"{{{self._A_NS}}}blip"
        embed_attr = f"{{{self._R_NS}}}embed"

        for blip in cell._tc.findall(".//" + blip_tag):
            rid = blip.get(embed_attr)
            if not rid:
                continue
            try:
                part = doc.part.rels[rid].target_part
                blob = part.blob
                ext = Path(part.partname).suffix.lower()
            except Exception as exc:
                log.warning(f"DOCX: rasm topilmadi (rId={rid}): {exc}")
                continue

            counter[0] += 1
            out_dir = settings.images_dir / doc_stem
            out_dir.mkdir(parents=True, exist_ok=True)
            base_name = f"t{table_idx}_row{row_idx}_img{counter[0]}"

            if ext in (".emf", ".wmf"):
                try:
                    from PIL import Image
                    image = Image.open(io.BytesIO(blob))
                    out_name = base_name + ".png"
                    image.convert("RGBA").save(out_dir / out_name)
                except Exception as exc:
                    log.warning(
                        f"DOCX: {ext} rasmni PNG ga o'girishda xatolik, "
                        f"asl format saqlanadi: {exc}"
                    )
                    out_name = base_name + ext
                    (out_dir / out_name).write_bytes(blob)
            else:
                out_name = base_name + (ext or ".png")
                (out_dir / out_name).write_bytes(blob)

            paths.append(f"images/{doc_stem}/{out_name}")

        return paths

    def _get_cell_texts_with_images(
        self, doc, row, doc_stem: str, table_idx: int, row_idx: int, counter: list[int]
    ) -> list[str]:
        """
        Like :meth:`_get_unique_cell_texts`, but appends a
        ``[[img:path]]`` token to each cell's text for every embedded image
        it contains — so a cell whose only content is an image is no
        longer empty.
        """
        seen: set[int] = set()
        texts: list[str] = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)

            text = self._get_cell_text(cell)
            images = self._extract_cell_images(doc, cell, doc_stem, table_idx, row_idx, counter)
            for img_path in images:
                token = build_image_token(img_path)
                text = (text + " " + token).strip() if text else token
            texts.append(text)
        return texts

    def _count_unique_columns(self, tbl) -> int:
        """Count unique (non-merged) columns by inspecting the first non-trivial row."""
        for row in tbl.rows:
            seen: set[int] = set()
            count = 0
            for cell in row.cells:
                tc_id = id(cell._tc)
                if tc_id not in seen:
                    seen.add(tc_id)
                    count += 1
            if count > 1:
                return count
        return 0

    def _is_table_columns_qa_format(self, doc) -> bool:
        """
        Return True when the document uses the multi-column table Q&A format.

        Detection criteria:
        - At least one table with 3+ unique columns (after merged-cell dedup)
        - 3+ rows have 2+ non-empty unique cells
        - No existing Classic markers (?, +, =) in body paragraphs
        """
        if not doc.tables:
            return False

        for para in doc.paragraphs[:30]:
            t = para.text.strip()
            if t and t[0] in "?+=":
                return False

        for tbl in doc.tables:
            if self._count_unique_columns(tbl) < 3:
                continue

            qa_rows = 0
            for row in tbl.rows[:30]:
                cells = self._get_unique_cell_texts(row)
                if sum(1 for c in cells if c) >= 2:
                    qa_rows += 1

            if qa_rows >= 3:
                return True

        return False

    def _extract_table_columns_qa(self, doc, path: Path) -> list[str]:
        """
        Extract multi-column table Q&A rows into Classic marker lines using TableParser.

        Iterates all tables, deduplicates merged cells per row, extracts any
        embedded images as ``[[img:path]]`` tokens, then delegates layout
        detection and line generation to TableParser.to_classic_lines().
        """
        from app.parser.table_parser import TableParser

        all_lines: list[str] = []
        doc_stem = self._sanitize_doc_stem(path.stem)
        img_counter = [0]

        for table_idx, tbl in enumerate(doc.tables):
            if self._count_unique_columns(tbl) < 2:
                continue

            rows: list[list[str]] = [
                self._get_cell_texts_with_images(doc, row, doc_stem, table_idx, row_idx, img_counter)
                for row_idx, row in enumerate(tbl.rows)
            ]
            lines = TableParser.to_classic_lines(rows)
            all_lines.extend(lines)

        if img_counter[0]:
            log.info(
                f"DOCX: {img_counter[0]} ta rasm ajratib olindi -> "
                f"{settings.images_dir / doc_stem}"
            )

        return all_lines
