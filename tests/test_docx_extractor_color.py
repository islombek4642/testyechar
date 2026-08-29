import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

from app.extractor.docx_extractor import DOCXExtractor


def _add_numpr(paragraph, num_id="1"):
    """Attach Word list-numbering metadata (w:numPr) directly to a
    paragraph's pPr, the way Word's own auto-numbered lists do -- the
    visible "1.", "2." is rendered from this, not stored as run text."""
    pPr = paragraph._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), num_id)
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)


def _make_docx(tmp_path):
    """A "-" (dash) marked question/option document where the correct
    option is colored red and bold, wrong options are colored a neutral
    dark blue -- mirrors real test banks that mark the answer by color
    instead of a highlight box or explicit +/# marker. The question line
    itself carries no marker at all (just plain text ending in "?")."""
    doc = docx.Document()
    doc.add_paragraph("Savol matni shu yerda?")
    doc.add_paragraph("- Noto'g'ri variant 1")
    p_correct = doc.add_paragraph()
    run = p_correct.add_run("- To'g'ri variant")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.bold = True
    p_wrong = doc.add_paragraph()
    run2 = p_wrong.add_run("- Noto'g'ri variant 2")
    run2.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    path = tmp_path / "red_option.docx"
    doc.save(str(path))
    return path


def test_red_text_option_gets_marked_correct_and_question_gets_prefix(tmp_path):
    path = _make_docx(tmp_path)
    result = DOCXExtractor().extract(path)
    text = result.full_text

    assert "? Savol matni shu yerda?" in text
    assert "+ To'g'ri variant" in text
    # The dash must be consumed by the correct-marker conversion, not left
    # sitting inside the option's text ("+ - To'g'ri variant" would be wrong).
    assert "+ - To'g'ri variant" not in text
    assert "- Noto'g'ri variant 1" in text
    assert "- Noto'g'ri variant 2" in text


def test_check_is_correct_option_detects_dominant_red(tmp_path):
    path = _make_docx(tmp_path)
    doc = docx.Document(str(path))
    ext = DOCXExtractor()
    paragraphs = list(ext._iter_all_paragraphs(doc))

    results = [ext._check_is_correct_option(p) for p in paragraphs]
    # paragraph order: question, wrong1, correct(red), wrong2
    assert results == [False, False, True, False]


def test_word_auto_numbered_question_gets_question_prefix(tmp_path):
    """A question paragraph carrying Word's own list-numbering metadata
    (numPr) -- rather than a literal "1." in its text, or being detected
    via the "next line is a dash option" fallback -- must still get the
    "?" prefix ClassicParser needs to recognize it as a question."""
    doc = docx.Document()
    q1 = doc.add_paragraph("Birinchi savol matni")
    _add_numpr(q1)
    doc.add_paragraph("- variant A")
    p_correct = doc.add_paragraph()
    run = p_correct.add_run("- variant B")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run.font.bold = True

    q2 = doc.add_paragraph("Ikkinchi savol matni")
    _add_numpr(q2)
    doc.add_paragraph("- variant C")
    doc.add_paragraph("- variant D")

    path = tmp_path / "numbered_questions.docx"
    doc.save(str(path))

    result = DOCXExtractor().extract(path)
    text = result.full_text

    assert "? Birinchi savol matni" in text
    assert "? Ikkinchi savol matni" in text
    assert "+ variant B" in text


def test_bold_with_no_color_marks_correct_option(tmp_path):
    """Some test banks mark the correct option by making it plain bold
    (default/automatic black text, no explicit color) against non-bold
    wrong options -- no highlight, no red/green, just weight."""
    doc = docx.Document()
    doc.add_paragraph("Savol matni?")
    doc.add_paragraph("- Noto'g'ri variant 1")
    p_correct = doc.add_paragraph()
    p_correct.add_run("- To'g'ri variant").bold = True
    doc.add_paragraph("- Noto'g'ri variant 2")

    path = tmp_path / "bold_only.docx"
    doc.save(str(path))

    result = DOCXExtractor().extract(path)
    assert "+ To'g'ri variant" in result.full_text


def test_bold_wrong_colored_option_not_mistaken_for_correct(tmp_path):
    """A WRONG option that happens to be bold while still carrying its
    ordinary (non-red/green) color -- bold used for unrelated emphasis,
    not as the correct-answer scheme -- must not be flagged correct."""
    doc = docx.Document()
    doc.add_paragraph("Savol matni?")
    p_bold_wrong = doc.add_paragraph()
    run = p_bold_wrong.add_run("- Noto'g'ri, lekin bold")
    run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
    run.bold = True
    p_correct = doc.add_paragraph()
    run2 = p_correct.add_run("- To'g'ri variant")
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    run2.bold = True
    doc.add_paragraph("- Noto'g'ri variant")

    path = tmp_path / "bold_wrong_colored.docx"
    doc.save(str(path))

    result = DOCXExtractor().extract(path)
    text = result.full_text
    assert "+ To'g'ri variant" in text
    assert "+ Noto'g'ri, lekin bold" not in text
    assert "- Noto'g'ri, lekin bold" in text


def test_hyperlink_wrapped_option_text_is_not_lost(tmp_path):
    """Word auto-links bare email addresses and URLs, nesting their run(s)
    one level deeper inside a <w:hyperlink> wrapper. Text extraction must
    still recover that text -- losing it is disastrous when the linked
    text is the correct answer itself (silently produces an empty option)."""
    doc = docx.Document()
    doc.add_paragraph("Savol matni?")
    doc.add_paragraph("- Oddiy variant")

    p_correct = doc.add_paragraph()
    # "- " as an ordinary direct run carrying the correct-answer color/bold
    # -- exactly how the real document is structured: the marker prefix is
    # a plain run, and Word's auto-link wraps only the linked text itself.
    prefix_run = p_correct.add_run("- ")
    prefix_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    prefix_run.bold = True
    # Manually build a <w:hyperlink> wrapping a run, mirroring what Word's
    # auto-link produces for a bare email address / URL.
    hyperlink = OxmlElement("w:hyperlink")
    run_elem = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = "dassa@umail.uz"
    run_elem.append(t)
    hyperlink.append(run_elem)
    p_correct._p.append(hyperlink)

    doc.add_paragraph("- Boshqa variant")

    path = tmp_path / "hyperlink_option.docx"
    doc.save(str(path))

    result = DOCXExtractor().extract(path)
    text = result.full_text
    assert "dassa@umail.uz" in text
    assert "+ dassa@umail.uz" in text
